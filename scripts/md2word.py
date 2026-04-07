#!/usr/bin/env python3
"""
md2word.py — 通用 Markdown → 专业 Word 文档转换工具

支持：标题(H1-H4)、表格、代码块、引用、有序/无序列表(含嵌套)、
     行内格式(加粗/斜体/行内代码/链接/删除线)、水平线、页码页脚

用法：
    python md2word.py <input.md>               # 输出同名 .docx
    python md2word.py <input.md> <output.docx> # 指定输出路径
    python md2word.py *.md                     # (bash glob) 批量转换
"""

import sys, re, os, argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════
#  配色方案
# ═══════════════════════════════════════════════
DARK_BLUE  = RGBColor(0x1F, 0x35, 0x64)
MID_BLUE   = RGBColor(0x2E, 0x54, 0x96)
SLATE_BLUE = RGBColor(0x44, 0x72, 0xC4)
BODY_COLOR = RGBColor(0x26, 0x26, 0x26)
CODE_COLOR = RGBColor(0x1E, 0x1E, 0x1E)
LINK_COLOR = RGBColor(0x00, 0x5C, 0xB8)

HEX_TBL_HEAD = "1F3864"
HEX_TBL_ALT  = "D9E2F3"
HEX_TBL_BORD = "2E5496"
HEX_CODE_BG  = "F5F5F5"
HEX_QUOT_BG  = "EDF2FA"
HEX_QUOT_LN  = "2E5496"


# ═══════════════════════════════════════════════
#  页头 / 页尾配置（在此修改公司信息、Logo、联系方式）
# ═══════════════════════════════════════════════

# ── 页头 ────────────────────────────────────────
ENABLE_HEADER         = False               # 是否启用页头
HEADER_LOGO_FILE      = 'realxen_logo.png'  # Logo 文件名（与脚本同目录）
HEADER_LOGO_HEIGHT_CM = 1.2                 # Logo 高度（cm），宽度自动等比
HEADER_DISTANCE_CM    = 1.0                 # 页眉距页面顶部距离（cm）
HEADER_BRAND_TEXT     = 'Realxen'           # 无 logo 时显示的品牌名
HEADER_BRAND_SUBTITLE = '  同信智维'         # 品牌名右侧副标题

# ── 页尾 ────────────────────────────────────────
ENABLE_FOOTER       = False               # 是否启用页尾
FOOTER_COMPANY_LINE = (
    '深圳市同信智维网络技术有限公司  '
    '地址：深圳市宝安区西乡街道办前进二路和流塘路交汇中粮锦云3栋1103-1106'
)
FOOTER_CONTACT_LINE = (
    '邮箱：info@realxen.com  电话：17620341860  '
    '微信：xmusmart  QQ：18755444  网站：www.realxen.com'
)
FOOTER_TAB_POS = '10546'  # 右对齐制表位 twips（A4 内容宽 18.6cm ≈ 10546 twips）


# ═══════════════════════════════════════════════
#  低级 XML 工具
# ═══════════════════════════════════════════════
def _set_cell_bg(cell, hex6: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)


def _set_table_borders(table, color=HEX_TBL_BORD):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        t = OxmlElement(f'w:{edge}')
        t.set(qn('w:val'),   'single')
        t.set(qn('w:sz'),    '4')
        t.set(qn('w:space'), '0')
        t.set(qn('w:color'), color)
        bdr.append(t)
    tblPr.append(bdr)


def _clear_table_borders(table):
    """移除表格所有边框，用于签字栏等无边框场景"""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        t = OxmlElement(f'w:{edge}')
        t.set(qn('w:val'),   'none')
        t.set(qn('w:sz'),    '0')
        t.set(qn('w:space'), '0')
        t.set(qn('w:color'), 'auto')
        bdr.append(t)
    tblPr.append(bdr)


def _set_spacing(para, before=0, after=0, line=None):
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'),  str(after))
    if line:
        sp.set(qn('w:line'),     str(line))
        sp.set(qn('w:lineRule'), 'auto')
    para._p.get_or_add_pPr().append(sp)


def _set_indent(para, left_twips: int, hanging=0):
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left_twips))
    if hanging:
        ind.set(qn('w:hanging'), str(hanging))
    para._p.get_or_add_pPr().append(ind)


def _set_para_shading(para, hex6: str):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    para._p.get_or_add_pPr().append(shd)


def _set_para_left_border(para, hex6: str, sz='24'):
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    sz)
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), hex6)
    pBdr.append(left)
    para._p.get_or_add_pPr().append(pBdr)


def _set_para_bottom_border(para, hex6: str, sz='6'):
    pBdr = OxmlElement('w:pBdr')
    b    = OxmlElement('w:bottom')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    sz)
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), hex6)
    pBdr.append(b)
    para._p.get_or_add_pPr().append(pBdr)


def _run_font(run, size=10.5, bold=False, italic=False,
              color=None, mono=False, strike=False, underline=False):
    if mono:
        run.font.name = 'Consolas'
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        except Exception:
            pass
    else:
        run.font.name = '微软雅黑'
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except Exception:
            pass
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.strike    = strike
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _run_inline_code_bg(run):
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  HEX_CODE_BG)
    rPr.append(shd)


# ═══════════════════════════════════════════════
#  页脚页码
# ═══════════════════════════════════════════════
def _add_footer(doc):
    footer = doc.sections[0].footer

    # ── 第一行：分隔线 + 公司名（左）+ 页码（右）──
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para, 60, 0)   # 段前 3pt 留出与分隔线的间距
    # 顶部细分隔线
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '4'); top.set(qn('w:color'), HEX_TBL_BORD)
    pBdr.append(top)
    para._p.get_or_add_pPr().append(pBdr)

    # 右侧页码制表位
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), FOOTER_TAB_POS)  # A4 可用宽度 18.6cm ≈ 10546 twips
    tabs.append(tab)
    pPr.append(tabs)

    def _run(para, text, size=8, bold=False, color=RGBColor(0x40, 0x40, 0x40)):
        r = para.add_run(text)
        r.font.name = '微软雅黑'
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color
        try: r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except Exception: pass
        return r

    _run(para, FOOTER_COMPANY_LINE, size=8, bold=False, color=RGBColor(0x40, 0x40, 0x40))

    tab_r = para.add_run()
    tab_r._element.append(OxmlElement('w:tab'))

    def _field(name):
        r = para.add_run()
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
        ins = OxmlElement('w:instrText'); ins.text = f' {name} '
        fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
        r._element.extend([fc1, ins, fc2])

    _field('PAGE')
    sep = para.add_run(' / ')
    sep.font.size = Pt(8)
    sep.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    _field('NUMPAGES')

    # ── 第二行：联系信息（单行，字号 7pt 确保不超行）──
    para2 = footer.add_paragraph()
    _set_spacing(para2, 0, 0)
    para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_text = FOOTER_CONTACT_LINE
    _run(para2, info_text, size=7, bold=False, color=RGBColor(0x88, 0x88, 0x88))


# ═══════════════════════════════════════════════
#  页头（Logo + 分隔线）
# ═══════════════════════════════════════════════
def _add_header(doc, logo_path: str = None):
    """在每页页头左侧添加 logo 图片（若有）或文字品牌标识，右侧留空，底部加分隔线。"""
    section = doc.sections[0]
    section.header_distance = Cm(HEADER_DISTANCE_CM)
    header = section.header

    # 清空默认空段落
    for p in header.paragraphs:
        p.clear()

    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_spacing(para, 0, 0)

    if logo_path and os.path.isfile(logo_path):
        # 嵌入 logo 图片，高度约 1.2cm，宽度按比例缩放
        run = para.add_run()
        run.add_picture(logo_path, height=Cm(HEADER_LOGO_HEIGHT_CM))
    else:
        # 无图片时，用样式化文字替代
        r1 = para.add_run(HEADER_BRAND_TEXT)
        r1.font.name   = 'Arial'
        r1.font.bold   = True
        r1.font.italic = True
        r1.font.size   = Pt(16)
        r1.font.color.rgb = RGBColor(0xD7, 0x19, 0x21)
        r2 = para.add_run(HEADER_BRAND_SUBTITLE)
        r2.font.name = '微软雅黑'
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 底部细分隔线
    _set_para_bottom_border(para, HEX_TBL_BORD, sz='4')


# ═══════════════════════════════════════════════
#  行内格式解析
# ═══════════════════════════════════════════════
_INLINE_RE = re.compile(
    r'(\*\*\*(?P<bi>[^*\n]+?)\*\*\*)'
    r'|(\*\*(?P<b>[^*\n]+?)\*\*)'
    r'|(__(?P<b2>[^_\n]+?)__)'
    r'|(\*(?P<i>[^*\n]+?)\*)'
    r'|(_(?P<i2>[^_\n]+?)_)'
    r'|(`(?P<code>[^`\n]+?)`)'
    r'|(!\[(?P<ia>[^\]\n]*?)\]\((?P<ip>[^)\n]+?)\))'
    r'|(\[(?P<lt>[^\]\n]+?)\]\((?P<lu>[^)\n]+?)\))'
    r'|(~~(?P<st>[^~\n]+?)~~)'
)


def _parse_inline(text: str) -> list:
    tokens, last = [], 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last:
            tokens.append(('text', text[last:m.start()]))
        g = m.groupdict()
        if   g['bi']:             tokens.append(('bi',     g['bi']))
        elif g['b']:              tokens.append(('bold',   g['b']))
        elif g['b2']:             tokens.append(('bold',   g['b2']))
        elif g['i']:              tokens.append(('ital',   g['i']))
        elif g['i2']:             tokens.append(('ital',   g['i2']))
        elif g['code']:           tokens.append(('code',   g['code']))
        elif g['ia'] is not None: tokens.append(('image',  g['ia'], g['ip']))
        elif g['lt']:             tokens.append(('link',   g['lt'], g['lu']))
        elif g['st']:             tokens.append(('strike', g['st']))
        last = m.end()
    if last < len(text):
        tokens.append(('text', text[last:]))
    return tokens


def _add_inline(para, text: str, size=10.5, color=None, bold=False):
    """将带行内格式的 Markdown 文本追加到段落"""
    for tok in _parse_inline(text):
        kind = tok[0]
        if kind == 'text':
            r = para.add_run(tok[1])
            _run_font(r, size=size, bold=bold, color=color)
        elif kind == 'bold':
            r = para.add_run(tok[1])
            _run_font(r, size=size, bold=True, color=color)
        elif kind == 'ital':
            r = para.add_run(tok[1])
            _run_font(r, size=size, italic=True, color=color)
        elif kind == 'bi':
            r = para.add_run(tok[1])
            _run_font(r, size=size, bold=True, italic=True, color=color)
        elif kind == 'code':
            r = para.add_run(f'\u202f{tok[1]}\u202f')
            _run_font(r, size=size - 0.5, mono=True, color=CODE_COLOR)
            _run_inline_code_bg(r)
        elif kind == 'link':
            r = para.add_run(tok[1])
            _run_font(r, size=size, color=LINK_COLOR, underline=True)
        elif kind == 'image':
            img_path = tok[2]
            base_dir = os.path.dirname(os.path.abspath(__file__))
            if not os.path.isabs(img_path):
                img_path = os.path.join(base_dir, img_path)
            if os.path.exists(img_path):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_path, height=Cm(3.0))
        elif kind == 'strike':
            r = para.add_run(tok[1])
            _run_font(r, size=size, color=RGBColor(0x80, 0x80, 0x80), strike=True)


# ═══════════════════════════════════════════════
#  块级元素构造器
# ═══════════════════════════════════════════════
def _doc_title(doc, text):
    """H1 首次出现 → 大标题（居中、22pt、深蓝）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before=280, after=160)
    r = p.add_run(text)
    _run_font(r, size=22, bold=True, color=DARK_BLUE)
    return p


def _toc_section_title(doc, text):
    """目录节标题 — 外观同 H2，但使用 Normal 样式段落，不被 Word TOC 字段收录"""
    p = doc.add_paragraph()
    _set_spacing(p, before=200, after=100)
    _set_para_bottom_border(p, '1F3864')
    r = p.add_run(text)
    _run_font(r, size=14, bold=True, color=DARK_BLUE)
    return p


def _heading(doc, text, level):
    """H2-H4 → Word 内置 Heading，可自动生成目录"""
    para = doc.add_heading(text, level=level)
    cfg  = {
        2: (14,   DARK_BLUE,  200, 100, '1F3864'),
        3: (12,   MID_BLUE,   160,  60,  None),
        4: (10.5, SLATE_BLUE, 100,  40,  None),
    }
    sz, col, bef, aft, bdr_col = cfg.get(level, (10.5, BODY_COLOR, 80, 30, None))
    if para.runs:
        r = para.runs[0]
        r.font.color.rgb = col
        r.font.size      = Pt(sz)
        r.font.bold      = True
        try:
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except Exception:
            pass
    _set_spacing(para, before=bef, after=aft)
    if bdr_col:
        _set_para_bottom_border(para, bdr_col)
    return para


def _paragraph(doc, text):
    p = doc.add_paragraph()
    _set_spacing(p, before=40, after=40, line=284)
    _add_inline(p, text, size=10.5, color=BODY_COLOR)
    return p


def _blockquote(doc, text):
    p = doc.add_paragraph()
    _set_spacing(p, before=80, after=80, line=280)
    _set_para_shading(p, HEX_QUOT_BG)
    _set_indent(p, left_twips=360)
    _set_para_left_border(p, HEX_QUOT_LN)
    _add_inline(p, text, size=10.5, color=MID_BLUE)
    return p


def _list_item(doc, text, level=0, ordered=False, num=1):
    BULLETS = ['•', '◦', '▪', '▫']
    p = doc.add_paragraph()
    _set_spacing(p, before=20, after=20, line=264)
    _set_indent(p, left_twips=360 + level * 360, hanging=240)
    br = p.add_run()
    if ordered:
        br.text = f'{num}.\t'
        _run_font(br, size=10.5, color=BODY_COLOR)
    else:
        br.text = BULLETS[level % len(BULLETS)] + '\t'
        _run_font(br, size=10.5, bold=True, color=MID_BLUE)
    _add_inline(p, text, size=10.5, color=BODY_COLOR)
    return p


def _code_block(doc, lines):
    for raw in lines:
        p = doc.add_paragraph()
        _set_spacing(p, 0, 0)
        _set_para_shading(p, HEX_CODE_BG)
        _set_indent(p, 240)
        r = p.add_run(raw or ' ')
        _run_font(r, size=9, mono=True, color=CODE_COLOR)
    sp = doc.add_paragraph()
    _set_spacing(sp, 0, 80)


def _hr(doc):
    p = doc.add_paragraph()
    _set_spacing(p, 80, 80)
    _set_para_bottom_border(p, 'BBBBBB')


def _table(doc, headers, rows):
    n   = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = 'Table Grid'
    _set_table_borders(tbl)

    # ── 按内容计算各列期望宽度 ──────────────────────────
    # A4 内容宽 = 21cm - 左右各 1.2cm = 18.6cm ≈ 10546 twips
    TOTAL_TWIPS    = 10546
    # 9.5pt 字号下：1 个"显示单位"(中文=2/英文=1) ≈ 95~100 twips
    TWIPS_PER_UNIT = 100
    # 单元格内边距（Word 默认左右各 108 twips）+ 小缓冲 = 300 twips
    CELL_PAD       = 300
    MIN_COL        = 640   # 最小列宽保底

    # 计算每列最长单行内容的"显示单位数"（<br> 分行后取最大值）
    col_units = []
    for ci in range(n):
        mx = _text_width(headers[ci].strip())
        for rd in rows:
            v = rd[ci].strip() if ci < len(rd) else ''
            for part in re.split(r'<br\s*/?>', v, flags=re.IGNORECASE):
                mx = max(mx, _text_width(part.strip()))
        col_units.append(max(mx, 2))

    # 期望宽度 = 最长行字宽 + 内边距缓冲
    desired = [max(u * TWIPS_PER_UNIT + CELL_PAD, MIN_COL) for u in col_units]
    total_d = sum(desired)

    if total_d <= TOTAL_TWIPS:
        # 总期望 ≤ 页宽：将剩余空间按期望比例分配给各列
        extra      = TOTAL_TWIPS - total_d
        twips_list = [d + int(extra * d / total_d) for d in desired]
    else:
        # 总期望 > 页宽：等比压缩，保底 MIN_COL
        twips_list = [max(int(TOTAL_TWIPS * d / total_d), MIN_COL) for d in desired]

    # 补齐舍入误差（加到最宽的列）
    diff = TOTAL_TWIPS - sum(twips_list)
    twips_list[twips_list.index(max(twips_list))] += diff

    # 固定布局 + 整表宽度
    tbl.autofit = False
    tbl_xml = tbl._tbl
    tblPr   = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    # 覆盖可能已存在的 tblW
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW_el = OxmlElement('w:tblW')
    tblW_el.set(qn('w:w'),    str(TOTAL_TWIPS))
    tblW_el.set(qn('w:type'), 'dxa')
    tblPr.append(tblW_el)

    # 重建 tblGrid（这才是 Word 固定布局下列宽的权威来源）
    for old in tbl_xml.findall(qn('w:tblGrid')):
        tbl_xml.remove(old)
    tblGrid = OxmlElement('w:tblGrid')
    for tw in twips_list:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(tw))
        tblGrid.append(gridCol)
    # tblGrid 必须紧跟在 tblPr 之后
    tblPr_idx = list(tbl_xml).index(tblPr)
    tbl_xml.insert(tblPr_idx + 1, tblGrid)

    # 每行每格同步写入 tcW（与 gridCol 保持一致）
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            _set_col_width(cell, twips_list[ci])

    # 表头
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, HEX_TBL_HEAD)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(para, 80, 80)
        r = para.add_run(h.strip())
        _run_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 数据行（隔行着色）
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = HEX_TBL_ALT if ri % 2 == 1 else 'FFFFFF'
        while len(row_data) < n:
            row_data.append('')
        for ci, ct in enumerate(row_data[:n]):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 支持 <br> 多行单元格
            parts = re.split(r'<br\s*/?>', ct.strip(), flags=re.IGNORECASE)
            for pi, part in enumerate(parts):
                if pi == 0:
                    para = cell.paragraphs[0]
                else:
                    para = cell.add_paragraph()
                _set_spacing(para, 40, 40)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_inline(para, part.strip(), size=9.5, color=BODY_COLOR)

    doc.add_paragraph()
    return tbl


# ═══════════════════════════════════════════════
#  表格列宽工具
# ═══════════════════════════════════════════════
def _text_width(s: str) -> int:
    """估算文本显示宽度：中文字符计2，ASCII 计1（已剥离 Markdown 标记）"""
    s = re.sub(r'\*+|_+|`|~~|!?\[[^\]]*\]\([^)]*\)', '', s)
    return sum(2 if '\u4e00' <= c <= '\u9fff' else 1 for c in s)


def _set_col_width(cell, twips: int):
    """精确设置单元格宽度（dxa twips）"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ═══════════════════════════════════════════════
#  Markdown 表格解析
# ═══════════════════════════════════════════════
def _is_sep(line: str) -> bool:
    s = line.strip()
    return bool(s) and s.startswith('|') and re.fullmatch(r'[\s|:\-]+', s)


def _sig_table(doc, left_lines, right_lines):
    """渲染双栏无边框签字区（甲方左、乙方右）"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    _clear_table_borders(tbl)
    col_w = Cm(9.0)
    for ci, side_lines in enumerate([left_lines, right_lines]):
        cell = tbl.rows[0].cells[ci]
        cell.width = col_w
        first = True
        for ln in side_lines:
            ln = ln.strip().rstrip('  ').rstrip()
            if not ln:
                continue
            if first:
                para = cell.paragraphs[0]
                first = False
            else:
                para = cell.add_paragraph()
            _set_spacing(para, 80, 80)
            _add_inline(para, ln, size=10.5, color=BODY_COLOR)
    doc.add_paragraph()


def _parse_table(lines, start):
    headers = [c for c in lines[start].strip().strip('|').split('|')]
    rows, i = [], start + 2
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append([c for c in lines[i].strip().strip('|').split('|')])
        i += 1
    return headers, rows, i


# ═══════════════════════════════════════════════
#  目录生成
# ═══════════════════════════════════════════════
_TOC_HEADING_NAMES = {'目录', 'table of contents', 'contents', 'toc', '目次'}


def _collect_headings(lines):
    """第一遍扫描 MD：收集所有标题（排除目录节标题本身），返回 (level, text) 列表"""
    headings = []
    in_code  = False
    first_h1 = False
    for line in lines:
        if re.match(r'^(`{3,}|~{3,})', line) and not in_code:
            in_code = True
            continue
        if in_code:
            if re.match(r'^(`{3,}|~{3,})\s*$', line):
                in_code = False
            continue
        hm = re.match(r'^(#{1,4})\s+(.*)', line)
        if not hm:
            continue
        level = len(hm.group(1))
        text  = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', hm.group(2)).strip()
        if text.strip().lower() in _TOC_HEADING_NAMES:
            continue
        if level == 1:
            if not first_h1:
                first_h1 = True
                continue   # 第一个 H1 是文档大标题，不纳入目录
            level = 2
        headings.append((level, text))
    return headings


def _toc_entries(doc, headings):
    """将收集到的标题渲染为目录条目段落（纯文本，兼容一切查看器）"""
    col_map = {2: DARK_BLUE, 3: MID_BLUE, 4: SLATE_BLUE}
    sz_map  = {2: 11.0,      3: 10.5,    4: 10.0}
    ind_map = {2: 0,         3: 360,     4: 720}
    for level, text in headings:
        p = doc.add_paragraph()
        _set_spacing(p, before=30, after=30)
        indent = ind_map.get(level, 0)
        if indent:
            _set_indent(p, left_twips=indent)
        r = p.add_run(text)
        _run_font(r, size=sz_map.get(level, 10.5),
                  color=col_map.get(level, BODY_COLOR))
    doc.add_paragraph()   # 目录后空行


# ═══════════════════════════════════════════════
#  文档初始化
# ═══════════════════════════════════════════════
def _setup_doc(doc, enable_header=ENABLE_HEADER, enable_footer=ENABLE_FOOTER):
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(1.2)
    sec.right_margin  = Cm(1.2)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(3.0)
    try:
        s = doc.styles['Normal']
        s.font.name = '微软雅黑'
        s.font.size = Pt(10.5)
        s._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    except Exception:
        pass
    # 自动查找同目录下的 logo 文件
    if enable_header:
        _logo = os.path.join(os.path.dirname(os.path.abspath(__file__)), HEADER_LOGO_FILE)
        _add_header(doc, logo_path=_logo)
    if enable_footer:
        _add_footer(doc)


# ═══════════════════════════════════════════════
#  主解析器
# ═══════════════════════════════════════════════
def convert(md_path: str, out_path: str, enable_header=ENABLE_HEADER, enable_footer=ENABLE_FOOTER):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    headings = _collect_headings(lines)   # 第一遍：收集文档标题，供目录节渲染

    doc = Document()
    _setup_doc(doc, enable_header=enable_header, enable_footer=enable_footer)

    i             = 0
    in_code       = False
    code_lines    = []
    first_h1      = False
    ol_counters   = {}   # level → current number
    in_toc_sec    = False   # 当前是否在「目录」标题下的列表区域
    toc_inserted  = False   # Word TOC 域是否已插入

    while i < len(lines):
        line = lines[i]

        # ── 代码块（支持 ``` 和 ~~~）─────────────────
        fence = re.match(r'^(`{3,}|~{3,})', line)
        if fence and not in_code:
            in_code    = True
            code_lines = []
            i += 1
            continue
        if in_code:
            if re.match(r'^(`{3,}|~{3,})\s*$', line):
                in_code = False
                _code_block(doc, code_lines)
            else:
                code_lines.append(line)
            i += 1
            continue

        # ── 水平线 ───────────────────────────────────
        if re.match(r'^(\*{3,}|-{3,}|_{3,})\s*$', line.strip()):
            _hr(doc)
            i += 1
            continue

        # ── 表格 ─────────────────────────────────────
        if line.strip().startswith('|') and i + 1 < len(lines) and _is_sep(lines[i + 1]):
            hdr, rows, ni = _parse_table(lines, i)
            _table(doc, hdr, rows)
            i = ni
            ol_counters.clear()
            continue
        if _is_sep(line):
            i += 1
            continue

        # ── 标题 ─────────────────────────────────────
        hm = re.match(r'^(#{1,4})\s+(.*)', line)
        if hm:
            level = len(hm.group(1))
            # 去掉标题中的 markdown 链接语法，保留文字
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', hm.group(2)).strip()
            # 检测是否是目录标题
            if text.strip().lower() in _TOC_HEADING_NAMES:
                in_toc_sec   = True
                toc_inserted = False
                # 用普通段落样式渲染，避免「目录」被 Word TOC 域收录
                _toc_section_title(doc, text)
            else:
                in_toc_sec = False   # 其他标题出现则退出 TOC 区域
                if level == 1 and not first_h1:
                    _doc_title(doc, text)
                    first_h1 = True
                elif level == 1:
                    _heading(doc, text, 2)   # 多余的 H1 → H2
                else:
                    _heading(doc, text, level)
            ol_counters.clear()
            i += 1
            continue

        # ── 签字栏（双栏无边框）────────────────────────
        if line.strip() == '<!-- sig:start -->':
            left_lines, right_lines = [], []
            i += 1
            while i < len(lines) and lines[i].strip() != '<!-- sig:split -->':
                left_lines.append(lines[i])
                i += 1
            i += 1  # 跳过 split
            while i < len(lines) and lines[i].strip() != '<!-- sig:end -->':
                right_lines.append(lines[i])
                i += 1
            i += 1  # 跳过 end
            _sig_table(doc, left_lines, right_lines)
            continue

        # ── 引用 ─────────────────────────────────────
        if line.startswith('>'):
            in_toc_sec = False
            text = re.sub(r'^>+\s?', '', line).strip()
            if text:
                _blockquote(doc, text)
            i += 1
            continue

        # ── 无序列表 ─────────────────────────────────
        ul = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if ul:
            if in_toc_sec:
                # TOC 区域内的列表 → 渲染目录条目段落（只渲染一次）
                if not toc_inserted:
                    _toc_entries(doc, headings)
                    toc_inserted = True
                i += 1
                continue
            lvl = len(ul.group(1)) // 2
            _list_item(doc, ul.group(2), level=lvl, ordered=False)
            ol_counters.clear()
            i += 1
            continue

        # ── 有序列表 ─────────────────────────────────
        ol = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if ol:
            if in_toc_sec:
                # TOC 区域内的列表 → 渲染目录条目段落（只渲染一次）
                if not toc_inserted:
                    _toc_entries(doc, headings)
                    toc_inserted = True
                i += 1
                continue
            lvl = len(ol.group(1)) // 2
            num = ol_counters.get(lvl, 0) + 1
            ol_counters[lvl] = num
            for k in list(ol_counters):
                if k > lvl:
                    del ol_counters[k]
            _list_item(doc, ol.group(3), level=lvl, ordered=True, num=num)
            i += 1
            continue

        # ── 空行 ─────────────────────────────────────
        if not line.strip():
            ol_counters.clear()
            i += 1
            continue

        # ── 普通段落 ─────────────────────────────────
        _paragraph(doc, line.strip())
        ol_counters.clear()
        i += 1

    doc.save(out_path)
    print(f'✅  {os.path.basename(md_path)}  →  {out_path}')


# ═══════════════════════════════════════════════
#  命令行入口
# ═══════════════════════════════════════════════
def main():
    ap = argparse.ArgumentParser(
        description='Markdown → 专业 Word 文档',
        epilog=(
            '示例：\n'
            '  python3 md2word.py README.md\n'
            '  python3 md2word.py a.md b.docx\n'
            '  python3 md2word.py -o out.docx a.md\n'
            '  python3 md2word.py *.md'
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument('args',    nargs='+', help='输入 .md 文件（可多个），最后一个参数若以 .docx 结尾则作为输出路径')
    ap.add_argument('-o', '--output',     help='指定输出路径（仅单文件时有效）')
    ap.add_argument('--header', action='store_true', default=ENABLE_HEADER,
                    help='启用页头（Logo + 品牌标识）')
    ap.add_argument('--footer', action='store_true', default=ENABLE_FOOTER,
                    help='启用页尾（公司信息 + 页码）')
    ap.add_argument('--no-header', dest='header', action='store_false',
                    help='禁用页头')
    ap.add_argument('--no-footer', dest='footer', action='store_false',
                    help='禁用页尾')
    parsed = ap.parse_args()

    positional = parsed.args

    # 自动识别：最后一个参数若以 .docx/.doc 结尾，视为输出路径
    if len(positional) >= 2 and positional[-1].lower().endswith(('.docx', '.doc')):
        inputs     = positional[:-1]
        out_single = positional[-1]
    else:
        inputs     = positional
        out_single = parsed.output  # 可能为 None

    if len(inputs) > 1 and out_single:
        print('⚠️  多个输入文件时忽略指定的输出路径，自动按输入文件名生成。', file=sys.stderr)
        out_single = None

    for md in inputs:
        if not os.path.isfile(md):
            print(f'❌  找不到文件：{md}', file=sys.stderr)
            continue
        # 输出路径：指定 > 与输入文件同目录同名.docx
        if out_single and len(inputs) == 1:
            # 相对路径：相对于当前工作目录解析（符合命令行直觉）
            out = os.path.abspath(out_single)
        else:
            out = os.path.splitext(os.path.abspath(md))[0] + '.docx'
        convert(md, out, enable_header=parsed.header, enable_footer=parsed.footer)


if __name__ == '__main__':
    main()
